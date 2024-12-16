#!/usr/bin/env python3
import sys
import os
import re
from docx import Document, opc
from PyQt6.QtCore import Qt, pyqtSignal
from PyQt6.QtWidgets import (QFileDialog, QApplication, QMainWindow,
                             QPushButton, QLabel, QLineEdit, QGridLayout,
                             QWidget, QTextEdit, QButtonGroup,QInputDialog, QDialog)
from PyQt6.QtGui import (QFont, QTextCursor, QAction,
                         QTextCharFormat, QKeySequence)
from functools import partial
import inputs_file
import openai
import random


# custom window subclass for custom prompts
class CustomWindow(QMainWindow):
    submitClicked = pyqtSignal(str)

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Make your own mad libs")
        self.error_text = ""
        self.custom_text_input = QLineEdit()
        self.custom_label = QLabel("Enter your custom prompt: ")
        self.custom_button = QPushButton("Done")
        self.custom_button.clicked.connect(self.close_window)

        layout = QGridLayout()
        layout.addWidget(self.custom_label, 0, 1, 1, 1)
        layout.addWidget(self.custom_text_input, 0, 2, 1, 22)
        layout.addWidget(self.custom_button, 1, 0, 10, 22)

        container = QWidget()
        container.setLayout(layout)
        # Set the central widget of the Window.
        self.setCentralWidget(container)

    def keyPressEvent(self, event):
        '''
        Function to look for the enter key press while
        filling in prompts to move to the next fill in the blank

        Keyword Arguments:
        event -- key press (looking for enter value: 16777220)
        '''
        if event.key() == 16777220:
            if self.custom_text_input.text() != "":
                self.close_window()
            else:
                pass

    def close_window(self):
        '''
        Function to close opened window and send input text back to main window
        Text is the custom prompt entered by the user
        If text is blank, nothing happens
        Activated by the Done button
        '''
        if self.custom_text_input.text() != "":
            self.submitClicked.emit(self.custom_text_input.text())
            self.close()
        else:
            pass


# Subclass QMainWindow to customize your application's main window
class MainWindow(QMainWindow):
    cursorMoved = pyqtSignal()

    def __init__(self):
        super().__init__()

        self.prompt_counter = 0
        self.added_prompts = []  # holds dict of prompts for cycling later
        self.added_prompts_dict = {}
        self.fill_in_blanks_mode = False
        self.final_pre_text = ""
        self.current_text = ""
        self.setWindowTitle("Make your own mad libs")
        self.buttons = []
        self.buttons_shortcut_label = []
        self.prompt_answers = {}

        # ENTER YOUR OPENAI KEY HERE
        openai.api_key = ''

        self.client = openai.OpenAI(
            # This is the default and can be omitted
            api_key=os.environ.get("OPENAI_API_KEY"),
        )

        # file menu options
        mainMenu = self.menuBar()
        # Save
        save_file_action = QAction("&Save File", self)
        save_file_action.setShortcut("Ctrl+S")
        save_file_action.setStatusTip('Save File')
        save_file_action.triggered.connect(self.file_save)
        # Open
        open_file_action = QAction("&Open File", self)
        open_file_action.setShortcut("Ctrl+O")
        open_file_action.setStatusTip('Open File')
        open_file_action.triggered.connect(self.file_open)
        # change background image
        change_image_action = QAction("Change Background", self)
        #change_image_action.setShortcut("Ctrl+B")
        change_image_action.setStatusTip('Change Background')
        change_image_action.triggered.connect(self.change_image)
        # Use ChatGPT
        chatgpt_action = QAction("Generate a mad lib", self)
        chatgpt_action.setStatusTip("Use ChatGPT")
        chatgpt_action.triggered.connect(self.ai_generate)
        # add menu options
        fileMenu = mainMenu.addMenu('&File')
        fileMenu.addAction(save_file_action)
        fileMenu.addAction(open_file_action)
        fileMenu.addAction(change_image_action)
        fileMenu.addAction(chatgpt_action)

        # add all the buttons dynamically based on the dictionary list
        self.button_list = inputs_file.inputs
        self.an_list = inputs_file.an_list
        self.prompt_group = QButtonGroup()  # group buttons together
        i = 0
        for key, value in self.button_list.items():
            self.buttons.append(QPushButton(value, self))
            self.buttons[i].clicked.connect(partial(self.add_a_prompt,
                                                    prompt=value))
            self.buttons[i].setShortcut(QKeySequence('Ctrl+' + key))
            # shortcut label
            self.buttons_shortcut_label.append(QLabel(key))
            self.buttons_shortcut_label[i].setStyleSheet("color:white;"
                                                "background-color:black")
            self.prompt_group.addButton(self.buttons[i])
            i = i + 1

        # Theme text and labels
        self.theme_text = QLineEdit()
        self.theme_label = QLabel("Theme:")
        self.theme_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.theme_label.setStyleSheet("color:white;background-color:black")
        self.theme_label.setFixedWidth(45)
        self.theme_label.setFixedHeight(20)

        # main text box and formatting for prompts
        self.full_text = QTextEdit()
        self.cursor = self.full_text.textCursor()
        self.full_text.cursorPositionChanged.connect(self.cursor_change)
        self.bold_fmt = QTextCharFormat()
        self.bold_fmt.setFontWeight(700)

        # labels for showing prompt count and prompt to fill in
        self.prompt_counter_label = QLabel("Number of Prompts: ")
        self.prompt_counter_label.setStyleSheet("color:white;"
                                                "background-color:black")

        # done button to begin the fill in the blank process
        self.done_button = QPushButton("Done")
        self.done_button.setEnabled(False)
        self.done_button.clicked.connect(self.start_fill_in_the_blank)
        # clear button to start over with blank text
        self.clear_button = QPushButton("Clear")
        self.clear_button.setEnabled(True)
        self.clear_button.clicked.connect(self.clear_all)
        # undo button for undoing prompts
        self.undo_button = QPushButton("Undo")
        self.undo_button.setEnabled(False)
        self.undo_button.setShortcut(QKeySequence('Ctrl+Z'))
        self.undo_button.clicked.connect(self.undo_prompt)

        # prompt editor buttons
        self.error_text = QLineEdit()
        self.error_text.setReadOnly(True)
        self.error_text.hide()

        # layout adds all items
        layout = QGridLayout()
        layout.addWidget(self.theme_label, 0, 1, 1, 1)
        layout.addWidget(self.theme_text, 0, 2, 1, 22)
        layout.addWidget(self.full_text, 1, 0, 10, 25)
        layout.addWidget(self.prompt_counter_label, 11, 0,
                         alignment=Qt.AlignmentFlag.AlignLeft)
        # add buttons based on the list
        for i in range(len(self.prompt_group.buttons())):
            if i < 10:
                layout.addWidget(self.prompt_group.buttons()[i], i + 1,
                                 27, alignment=Qt.AlignmentFlag.AlignLeft)
                layout.addWidget(self.buttons_shortcut_label[i], i + 1,
                                 26, alignment=Qt.AlignmentFlag.AlignRight)
            else:
                layout.addWidget(self.prompt_group.buttons()[i], i - 9,
                                 29, alignment=Qt.AlignmentFlag.AlignLeft)
                layout.addWidget(self.buttons_shortcut_label[i], i - 9,
                                 28, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.undo_button, 11, 23, 1, 1)
        layout.addWidget(self.clear_button,12 ,23, 1, 1)
        layout.addWidget(self.done_button, 12, 0, 1, 1)
        layout.addWidget(self.error_text, 14, 0, 1, 5)
        #layout.addWidget(self.undo_button, 12, 16, 1, 1)

        # Set the central widget of the Window.
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    ###################################################
    # Function definitions
    ###################################################

    def cursor_change(self):
        '''Function to track the cursor in case of undo events'''
        self.cursor.setPosition(self.full_text.textCursor().position())
        self.full_text.setTextCursor(self.cursor)

    def file_save(self):
        '''
        Function to save file. Currently saves as docx file with no formatting
        '''
        document = Document()
        document.add_heading(self.theme_text.text(), level=1)
        document.add_paragraph(self.full_text.toPlainText())
        document.add_paragraph()
        document.add_heading("Do it again!")
        document.add_paragraph(self.final_pre_text)
        document.add_paragraph(str(self.added_prompts_dict))
        name = QFileDialog.getSaveFileName(self, 'Save File')
        if name[0]:
            document.save(name[0] + '.docx')
        else:
            self.close()

    def file_open(self):
        '''
        Function to open a file. Used to re-do saved Mad Libs.
        '''
        name = QFileDialog.getOpenFileName(self, 'Open File',filter="Text files (*.docx)")
        try:
            document = Document(name[0])
            text_to_grab = False
            self.theme_text.setText(document.paragraphs[0].text) #  set theme
            # loop through paragraphs and find do it over text
            for paragraph in document.paragraphs:
                if paragraph.text == "":
                    pass
                elif paragraph.text.upper() == "DO IT AGAIN!":
                    # next loop contains text to grab
                    text_to_grab = True
                elif paragraph.text[0] == "{":
                    text_to_grab = False
                    pass
                elif text_to_grab == True:
                    self.current_text = paragraph.text
        except opc.exceptions.PackageNotFoundError:
            pass
        # convert text to sorted dictionary
        try:
            self.added_prompts = re.findall(r'\[\s*([a-zA-Z\s]+[0-9]*)\s*\]', self.current_text)
            for i, prompt in enumerate(self.added_prompts):
                self.added_prompts_dict[i] = prompt
            self.full_text.setText(self.current_text)
            self.theme_text.setReadOnly(False)
            self.full_text.setReadOnly(False)
            self.done_button.setEnabled(True)
            self.undo_button.setEnabled(True)
            self.prompt_counter_label_update()

            for button in self.prompt_group.buttons():
                button.setEnabled(True)
        # error if application is already open
        except:
            self.full_text.setText("INVALID CONFIG. MAKE A NEW MAD LIB YA GOOF")

    def change_image(self):
        '''
        Function to allow user to change the background image based on
        selecting a file
        '''
        image_path, _ = QFileDialog.getOpenFileName()
        if image_path:
            image_name = os.path.basename(image_path)
            self.setStyleSheet("MainWindow{ border-image:url("
                               + "backgrounds/" + image_name +
                               "); background-repeat: no-repeat;"
                               "background-position: center}")
            
    def ai_generate(self):
        '''
        Function to use ChatGPT to generate a mad lib based on the theme and number of prompts input by the user. Attempts to generate a one-paragraph mad lib based on the 
        chosen theme. Checks for the correct number of prompts and quits if it is unable to generate it with the correct number of prompts.
        Currently uses the GPT4.0 model.
        '''
        isCanceled = False
        retries = 1
        max_retries = 3
        # no current checking for good input. Uses defaults for blank answers.
        theme_madlib, done1 = QInputDialog.getText(self, 'Theme Selection', 'Enter the theme to generate:')
        if done1:
            if not theme_madlib:
                # picking a default of winter for now
                theme_madlib = "Winter"
            number_of_blanks, done2 = QInputDialog.getText(self, 'Prompt Number', 'Enter how many prompts should be generated:')
            if done2:
                if not number_of_blanks:
                    # choose a random number of prompts between 1 and 10 if none given
                    number_of_blanks = str(random.randint(1, 10))

                prompt = f"""Generate a one-paragraph Mad Lib based on the theme of {theme_madlib} using **exactly** {number_of_blanks} placeholders. 
                        The total number of placeholders in the paragraph **must be exactly {number_of_blanks}**. Do not generate any more or any fewer placeholders.
                        Here are the allowed prompt types: Verb Past Tense, Plural Noun, Silly word, Adjective, Body Part, Animal, Verb, Verb ending in ing, Adverb, Adjective, Number, Colour, Type of Liquid, Type of Food, Place, Celebrity, Exclamation, Person in Room, Part of the Body, Part of the Body Plural.
                        **Instructions**:
                        1. Include exactly {number_of_blanks} placeholders in the paragraph.
                        2. Mark each fill-in-the-blank spot in square brackets with ascending numbers (e.g., [adjective1], [verb1], [noun1]).
                        3. If duplicates are needed, give each a unique number (e.g., [verb1], [verb2]).
                        4. Count the number of placeholders in your response to ensure the total is *exactly* {number_of_blanks}. If the count is incorrect, regenerate the paragraph.

                        Do not list the placeholders separately; only integrate them into the paragraph. Double-check the output before responding to ensure the correct number of placeholders."""
                
                # try to generate the mad lib. If the incorrect number of prompts is used, try again up to a max of 3 tries before just giving up. Sometimes the model is dumb.
                while retries <= max_retries:
                    try:
                        chat_completion = self.client.chat.completions.create(
                        messages=[{"role": "user","content": prompt}],
                        temperature=0.2,
                        max_tokens=500,
                        model="gpt-4",
                        )
                        # Extract the response content
                        response_content = chat_completion.choices[0].message.content

                        # regex to remove prompts and paragraph
                        mad_lib_paragraph = re.sub(r'\n.*', '', response_content).strip()
                        self.final_pre_text = mad_lib_paragraph
                        prompt_matches = re.findall(r'\[\s*([a-zA-Z\s]+[0-9]*)\s*\]', response_content)
    
                        if len(prompt_matches) == int(number_of_blanks):
                            isCanceled = False
                            break
                        else:
                            retries += 1
                    except Exception as e:
                        self.error_text.setText(f"An error occurred: {e}")
                        self.error_text.show()
                        break

                if retries > max_retries:
                    self.full_text.setReadOnly(False)
                    self.full_text.setText("Failed to generate a mad lib correctly. Try again.")
                    self.close()
            else:
                # clicked cancel on the second prompt.
                isCanceled = True
            
            if isCanceled:
                self.close()    
            # fill dictionary with prompts
            else:
                prompt_dict = {}
                for i, prompt in enumerate(prompt_matches):
                    prompt_dict[i] = prompt

                mad_lib_filled = mad_lib_paragraph.format(**prompt_dict)
                self.theme_text.setText(theme_madlib)
                self.theme_text.setReadOnly(True)
                self.full_text.setText(mad_lib_filled)
                self.full_text.setReadOnly(True)
                self.done_button.setEnabled(False)
                self.undo_button.setEnabled(False)
                self.prompt_counter_label.setText("Number of Prompts: " + number_of_blanks)

                for button in self.prompt_group.buttons():
                    button.setEnabled(False)
                # loop through prompts and get responses
                prompt_responses = {}
                for key, value in prompt_dict.items():
                    if isCanceled:
                        self.close()
                    else:
                        prompt_answer, done3 = QInputDialog.getText(self, 'Fill-in-the-blanks!', 'Enter a(n) ' + str(value[:-1]) + ':')
                        if done3:
                            prompt_responses[value] = prompt_answer
                            isCanceled = False
                        else:
                            isCanceled = True

                # replace prompts with answers
                for key in prompt_responses:
                    mad_lib_paragraph = mad_lib_paragraph.replace(key, prompt_responses[key])
                
                self.theme_text.setReadOnly(False)
                self.full_text.setText(mad_lib_paragraph)
        else:
            # clicked cancel on first dialog
            self.close()

    def custom_prompt_window(self):
        '''
        Function to open a new window when the custom prompt is selected
        The window allows the user to fill in a custom prompt
        '''
        self.sub_window = CustomWindow()
        self.sub_window.submitClicked.connect(self.on_sub_window_confirm)
        self.sub_window.show()

    def on_sub_window_confirm(self, prompt):
        '''
        Function to fill in custom prompt received from custom_prompt_window

        Keyword Arguments
        prompt -- custom prompt (str)
        '''
        self.add_a_prompt(prompt)

    ###################################################
    # Helper Function definitions
    ###################################################

    def prompt_counter_label_update(self):
        '''Update prompt label with new count'''
        self.prompt_counter_label.setText("Number of Prompts: "
                                          + str(len(self.added_prompts)))
        
    def clear_all(self):
        self.full_text.setText("")
        self.full_text.setReadOnly(False)
        self.theme_text.setText("")
        self.theme_text.setReadOnly(False)
        self.done_button.setEnabled(True)
        self.undo_button.setEnabled(True)

        for button in self.prompt_group.buttons():
            button.setEnabled(True)
        
        self.added_prompts = []
        self.added_prompts_dict = {}
        self.prompt_counter_label_update()


    ###################################################
    # Button Function definitions
    ###################################################

    def add_a_prompt(self, prompt):
        '''
        Function to add a prompt which is based on the button clicked
        Activated by a prompt button

        Keyword arguments:
        prompt -- text of button which was pressed (str)
        '''
        if not prompt:
            # skip if no text in the prompt. Usually from custom prompts
            pass
        elif prompt.upper() == "CUSTOM":
            self.custom_prompt_window()
        else:
            # first one added then check if any were deleted
            if len(self.added_prompts) == 0:
                self.added_prompts.append(prompt + str(len(self.added_prompts) + 1))
                self.cursor.insertText("[" + prompt + str(len(self.added_prompts)) + "]")
                prompt_matches = re.findall(r'\[\s*([a-zA-Z\s]+[0-9]*)\s*\]', self.full_text.toPlainText())
                self.added_prompts = prompt_matches
            else:
                # if one prompt already was put in, check that any weren't deleted then add new one
                prompt_matches = re.findall(r'\[\s*([a-zA-Z\s]+[0-9]*)\s*\]', self.full_text.toPlainText())
                self.added_prompts = prompt_matches
                self.added_prompts.append(prompt + str(len(self.added_prompts) + 1))
                self.cursor.insertText("[" + prompt + str(len(self.added_prompts)) + "]")

            self.prompt_counter_label_update()
            self.full_text.setFocus()
            # show undo button after first prompt is entered
            if self.undo_button.isEnabled() is False:
                self.undo_button.setEnabled(True)
            if self.done_button.isEnabled() is False:
                self.done_button.setEnabled(True)

    def undo_prompt(self):
        '''
        Function to remove the last prompt entered in
        Separate from undoing the last fill-in-the-blank
        Activated by the Undo button
        '''
        if self.fill_in_blanks_mode:
            pass
        else:
            self.current_text = self.full_text.toPlainText()
            prompt_to_remove = self.added_prompts[-1]
            prompt_loc = self.current_text.find(prompt_to_remove)
            if prompt_loc >= 0:
            # move cursor to prompt location, then select and remove the text
                self.cursor.setPosition(prompt_loc - 1,
                                        QTextCursor.MoveMode.MoveAnchor)
                self.cursor.movePosition(QTextCursor.MoveOperation.Right,
                                        QTextCursor.MoveMode.KeepAnchor,
                                        len(prompt_to_remove) + 2)
                self.cursor.removeSelectedText()

                # put cursor back to where the prompt was removed
                self.full_text.setTextCursor(self.cursor)
                self.full_text.setFocus()

                prompt_matches = re.findall(r'\[\s*([a-zA-Z\s]+[0-9]*)\s*\]', self.full_text.toPlainText())
                self.added_prompts = prompt_matches

                # clean up behind the scenes
                self.prompt_counter_label_update()
                if len(self.added_prompts) == 0:
                    self.undo_button.setEnabled(False)
                    self.done_button.setEnabled(False)
            else:
                pass

    def start_fill_in_the_blank(self):
        '''
        Function to start filling in the blank prompt and then update the
        text after clicking next.
        Activated by the Done button
        '''
        self.prompt_counter_label_update()
        isCanceled = False
        if len(self.added_prompts) == 0:
            # error label
            self.error_text.show()
            self.error_text.setText("You didn't write any fill-in-the-blanks!")
            pass
        else:
            self.error_text.hide()
            self.final_pre_text = self.full_text.toPlainText()
            self.current_text = self.full_text.toPlainText()
            self.theme_text.setReadOnly(True)
            self.full_text.setReadOnly(True)
            self.done_button.setEnabled(False)
            self.undo_button.setEnabled(False)

            for button in self.prompt_group.buttons():
                button.setEnabled(False)
            # loop through prompts and get responses

            for i, prompt in enumerate(self.added_prompts):
                self.added_prompts_dict[i] = prompt

            prompt_responses = {}
            for key, value in self.added_prompts_dict.items():
                if isCanceled:
                    self.close()
                else:
                    prompt_answer, done3 = QInputDialog.getText(self, 'Fill-in-the-blanks!', 'Enter a(n) ' + str(value[:-1]) + ':')
                    if done3:
                        prompt_responses[value] = prompt_answer
                        isCanceled = False
                    else:
                        isCanceled = True

            # replace prompts with answers
            for key in prompt_responses:
                self.current_text = self.current_text.replace(key, prompt_responses[key])
            
            self.theme_text.setReadOnly(False)
            self.full_text.setReadOnly(False)
            self.full_text.setText(self.current_text)

stylesheet_main = """
        MainWindow {
        background-image: url("backgrounds/silly.png") repeat;
    }
"""

app = QApplication(sys.argv)
app.setStyleSheet(stylesheet_main)  # background image
window = MainWindow()
window.show()

app.exec()
